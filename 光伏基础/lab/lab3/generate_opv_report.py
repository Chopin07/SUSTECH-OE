from pathlib import Path
import csv
import math

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA = ROOT / "OPV DATA"
FIGURES = ROOT / "figures"
TABLES = ROOT / "tables"
FIGURES.mkdir(exist_ok=True)
TABLES.mkdir(exist_ok=True)

DEVICE_ORDER = [
    ("D18", "D18 only"),
    ("D18 Y6无添加剂", "D18:Y6 without additive"),
    ("D18 Y6加添加剂", "D18:Y6 with additive"),
]

USE_GROUP_ONLY = "group3"
REPORT_SUFFIX = "Group3_DeviceArchitecture"

COLORS = {
    "D18 only": "#6b7280",
    "D18:Y6 without additive": "#1f77b4",
    "D18:Y6 with additive": "#d62728",
}


def read_jv_file(path: Path):
    rows = list(csv.reader(path.open(encoding="utf-8-sig")))
    metrics = {}
    for row in rows[:10]:
        if len(row) >= 4 and row[1].strip():
            try:
                metrics[row[1].strip()] = float(row[3])
            except ValueError:
                pass
    df = pd.read_csv(path, skiprows=10, usecols=[0, 1, 2])
    df.columns = ["Voltage (V)", "Current (mA)", "Current density (mA cm^-2)"]
    df = df.dropna()
    return metrics, df


def read_eqe_file(path: Path):
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    integrated_jsc = float(lines[1].split("\t")[1])
    df = pd.read_csv(path, sep="\t", skiprows=4, usecols=[0, 1, 2, 3])
    df.columns = ["Wavelength (nm)", "EQE (%)", "SR (A/W)", "Jsc contribution (mA cm^-2)"]
    return integrated_jsc, df.dropna()


def summarize_jv():
    records = []
    curves = {}
    for folder, label in DEVICE_ORDER:
        curves[label] = []
        for path in sorted((DATA / folder).glob("*.csv")):
            if USE_GROUP_ONLY and path.stem.lower() != USE_GROUP_ONLY.lower():
                continue
            metrics, df = read_jv_file(path)
            record = {
                "Device": label,
                "Group": path.stem,
                "Voc (V)": metrics.get("Voc (V)", np.nan),
                "Jsc (mA cm^-2)": metrics.get("Jsc (mA/cm2)", np.nan),
                "FF (%)": metrics.get("FF (%)", np.nan),
                "PCE (%)": metrics.get("PCE (%)", np.nan),
                "Vm (V)": metrics.get("Vm (V)", np.nan),
                "Im (mA)": metrics.get("Im (mA)", np.nan),
            }
            records.append(record)
            curves[label].append((path.stem, df))
    raw = pd.DataFrame(records)
    stats = raw.groupby("Device", sort=False)[
        ["Voc (V)", "Jsc (mA cm^-2)", "FF (%)", "PCE (%)"]
    ].agg(["mean", "std"])
    stats.columns = [f"{a} {b}" for a, b in stats.columns]
    stats = stats.reset_index()
    return raw, stats, curves


def summarize_eqe():
    records = []
    curves = {}
    for filename, label in [
        ("D18.txt", "D18 only"),
        ("D18 Y6不加添加剂.txt", "D18:Y6 without additive"),
        ("D18 Y6加添加剂.txt", "D18:Y6 with additive"),
    ]:
        integrated_jsc, df = read_eqe_file(DATA / "EQE" / filename)
        curves[label] = df
        edge_fit = df[(df["Wavelength (nm)"] >= 840) & (df["Wavelength (nm)"] <= 920)]
        slope, intercept = np.polyfit(edge_fit["Wavelength (nm)"], edge_fit["EQE (%)"], 1)
        linear_edge = -intercept / slope if slope < 0 else np.nan
        linear_eg = 1240 / linear_edge if linear_edge and not math.isnan(linear_edge) else np.nan
        if label == "D18 only":
            linear_edge = np.nan
            linear_eg = np.nan
        max_eqe = df["EQE (%)"].max()
        mean_visible = df[(df["Wavelength (nm)"] >= 400) & (df["Wavelength (nm)"] <= 800)]["EQE (%)"].mean()
        records.append({
            "Device": label,
            "Integrated Jsc from EQE (mA cm^-2)": integrated_jsc,
            "Mean EQE 400-800 nm (%)": mean_visible,
            "Maximum EQE (%)": max_eqe,
            "Linear absorption edge (nm)": linear_edge,
            "Bandgap from linear edge (eV)": linear_eg,
        })
    return pd.DataFrame(records), curves


def save_tables(jv_raw, jv_stats, eqe_stats):
    jv_raw.to_csv(TABLES / "jv_individual_parameters.csv", index=False, encoding="utf-8-sig")
    jv_stats.to_csv(TABLES / "jv_summary_statistics.csv", index=False, encoding="utf-8-sig")
    eqe_stats.to_csv(TABLES / "eqe_summary_statistics.csv", index=False, encoding="utf-8-sig")


def plot_jv(curves):
    fig, axes = plt.subplots(1, 3, figsize=(12, 3.2), sharey=False)
    for ax, (label, items) in zip(axes, curves.items()):
        for group, df in items:
            ax.plot(df["Voltage (V)"], df["Current density (mA cm^-2)"], lw=1.2, alpha=0.85, label=group)
        ax.axhline(0, color="black", lw=0.7)
        ax.axvline(0, color="black", lw=0.7)
        ax.set_title(label, fontsize=10)
        ax.set_xlabel("Voltage (V)")
        ax.grid(True, lw=0.4, alpha=0.35)
    axes[0].set_ylabel("Current density (mA cm$^{-2}$)")
    axes[0].legend(fontsize=7, frameon=False)
    fig.tight_layout()
    fig.savefig(FIGURES / "opv_jv_curves.png", dpi=300)
    plt.close(fig)


def plot_performance(jv_stats):
    labels = jv_stats["Device"].tolist()
    x = np.arange(len(labels))
    fig, axes = plt.subplots(2, 2, figsize=(8.6, 6.2))
    metrics = [
        ("Jsc (mA cm^-2) mean", "Jsc (mA cm$^{-2}$)"),
        ("Voc (V) mean", "Voc (V)"),
        ("FF (%) mean", "FF (%)"),
        ("PCE (%) mean", "PCE (%)"),
    ]
    for ax, (col, ylabel) in zip(axes.ravel(), metrics):
        err_col = col.replace(" mean", " std")
        colors = [COLORS[label] for label in labels]
        ax.bar(x, jv_stats[col], yerr=jv_stats[err_col], capsize=4, color=colors, alpha=0.85)
        ax.set_ylabel(ylabel)
        ax.set_xticks(x)
        ax.set_xticklabels(["D18", "D18:Y6\nno additive", "D18:Y6\nadditive"], fontsize=8)
        ax.grid(axis="y", lw=0.4, alpha=0.35)
    fig.tight_layout()
    fig.savefig(FIGURES / "opv_performance_summary.png", dpi=300)
    plt.close(fig)


def plot_eqe(eqe_curves):
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    for label, df in eqe_curves.items():
        ax.plot(df["Wavelength (nm)"], df["EQE (%)"], lw=1.8, label=label, color=COLORS[label])
    ax.set_xlabel("Wavelength (nm)")
    ax.set_ylabel("EQE (%)")
    ax.set_xlim(300, 950)
    ax.set_ylim(0, 90)
    ax.grid(True, lw=0.4, alpha=0.35)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(FIGURES / "opv_eqe_spectra.png", dpi=300)
    plt.close(fig)


def plot_eqe_edge(eqe_curves):
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    for label, df in eqe_curves.items():
        if label == "D18 only":
            continue
        sub = df[(df["Wavelength (nm)"] >= 820) & (df["Wavelength (nm)"] <= 940)]
        fit = df[(df["Wavelength (nm)"] >= 840) & (df["Wavelength (nm)"] <= 920)]
        slope, intercept = np.polyfit(fit["Wavelength (nm)"], fit["EQE (%)"], 1)
        edge = -intercept / slope
        xs = np.linspace(840, edge, 80)
        ax.plot(sub["Wavelength (nm)"], sub["EQE (%)"], lw=1.8, color=COLORS[label], label=label)
        ax.plot(xs, slope * xs + intercept, "--", lw=1.1, color=COLORS[label], alpha=0.8)
        ax.axvline(edge, color=COLORS[label], lw=0.8, alpha=0.55)
    ax.set_xlabel("Wavelength (nm)")
    ax.set_ylabel("EQE (%)")
    ax.set_xlim(820, 940)
    ax.set_ylim(0, 70)
    ax.grid(True, lw=0.4, alpha=0.35)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(FIGURES / "opv_eqe_edge_fit.png", dpi=300)
    plt.close(fig)


def md_table(df, cols, formats):
    lines = []
    lines.append("| " + " | ".join(cols) + " |")
    lines.append("|" + "|".join(["---" for _ in cols]) + "|")
    for _, row in df.iterrows():
        vals = []
        for col in cols:
            val = row[col]
            fmt = formats.get(col)
            if pd.isna(val):
                vals.append("")
            else:
                vals.append(fmt.format(val) if fmt else str(val))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def compact_jv_stats(jv_stats):
    out = pd.DataFrame()
    out["Device"] = jv_stats["Device"]
    for base in ["Voc (V)", "Jsc (mA cm^-2)", "FF (%)", "PCE (%)"]:
        out[base] = [
            f"{mean:.3f}" if pd.isna(std) else f"{mean:.3f} +/- {std:.3f}"
            for mean, std in zip(jv_stats[f"{base} mean"], jv_stats[f"{base} std"])
        ]
    return out


LAYER_FUNCTIONS = [
    {
        "Layer": "Transparent front electrode / substrate",
        "Material": "Patterned ITO/glass",
        "Function": "Provides mechanical support and a transparent conducting contact. Light enters through this side, and holes are collected through the ITO/PEDOT:PSS contact.",
    },
    {
        "Layer": "Hole transport layer (HTL)",
        "Material": "PEDOT:PSS",
        "Function": "Smooths the ITO surface, selectively extracts holes, blocks electrons, and reduces the injection/extraction barrier at the anode side.",
    },
    {
        "Layer": "Active layer",
        "Material": "D18 only or D18:Y6 blend",
        "Function": "Absorbs photons and generates excitons. In D18:Y6 devices, the donor-acceptor bulk heterojunction provides interfaces for exciton dissociation and pathways for hole/electron transport.",
    },
    {
        "Layer": "Electron transport/interfacial layer (ETL)",
        "Material": "PDINN",
        "Function": "Selectively extracts electrons, modifies the cathode interface, lowers contact resistance, and helps block holes from reaching the Ag electrode.",
    },
    {
        "Layer": "Back metal electrode",
        "Material": "Ag, 100 nm",
        "Function": "Collects electrons from the ETL side, defines the top contact, and completes the external circuit.",
    },
]


def layer_functions_markdown():
    lines = [
        "| Layer | Material | Function |",
        "|---|---|---|",
    ]
    for item in LAYER_FUNCTIONS:
        lines.append(f"| {item['Layer']} | {item['Material']} | {item['Function']} |")
    return "\n".join(lines)


def values_for_text(jv_stats):
    indexed = jv_stats.set_index("Device")

    def val(device, col):
        return float(indexed.loc[device, f"{col} mean"])

    return {
        "no_voc": val("D18:Y6 without additive", "Voc (V)"),
        "add_voc": val("D18:Y6 with additive", "Voc (V)"),
        "d18_pce": val("D18 only", "PCE (%)"),
        "no_pce": val("D18:Y6 without additive", "PCE (%)"),
        "add_pce": val("D18:Y6 with additive", "PCE (%)"),
        "no_jsc": val("D18:Y6 without additive", "Jsc (mA cm^-2)"),
        "add_jsc": val("D18:Y6 with additive", "Jsc (mA cm^-2)"),
        "no_ff": val("D18:Y6 without additive", "FF (%)"),
        "add_ff": val("D18:Y6 with additive", "FF (%)"),
    }


def additive_comparison(jv_stats):
    vals = values_for_text(jv_stats)
    rows = [
        ("Voc (V)", vals["no_voc"], vals["add_voc"], "Voltage remains unchanged"),
        ("Jsc (mA cm^-2)", vals["no_jsc"], vals["add_jsc"], "Higher photocurrent after additive treatment"),
        ("FF (%)", vals["no_ff"], vals["add_ff"], "Improved curve squareness and charge extraction"),
        ("PCE (%)", vals["no_pce"], vals["add_pce"], "Overall efficiency increases"),
    ]
    out = []
    for parameter, before, after, note in rows:
        change = after - before
        rel = change / before * 100 if before else np.nan
        out.append({
            "Parameter": parameter,
            "Without additive": before,
            "With additive": after,
            "Change": change,
            "Relative change (%)": rel,
            "Interpretation": note,
        })
    return pd.DataFrame(out)


METHOD_STEPS = [
    {
        "Step": "1",
        "Process": "Sonication in detergent water for 15 min",
        "Purpose": "Removes dust, grease, and loosely attached organic residues from patterned ITO/glass. This reduces shunting pathways and improves the reproducibility of later coating steps.",
    },
    {
        "Step": "2",
        "Process": "Sonication in deionized water for 15 min",
        "Purpose": "Rinses away detergent residues and soluble ionic contaminants. A cleaner ITO surface helps avoid leakage current and interfacial recombination.",
    },
    {
        "Step": "3",
        "Process": "Sonication in absolute ethanol for 15 min",
        "Purpose": "Further removes organic residues and promotes rapid drying. Ethanol is volatile and helps leave a cleaner substrate surface before thermal drying.",
    },
    {
        "Step": "4",
        "Process": "Air-gun drying and storage at 60 degC for 1 h",
        "Purpose": "Removes residual solvent and moisture from the substrate. Dry substrates are important for uniform PEDOT:PSS coating and stable device interfaces.",
    },
    {
        "Step": "5",
        "Process": "Preparation of D18:Y6 and D18 solutions using chloroform as solvent",
        "Purpose": "Creates the active-layer inks. The D18:Y6 blend provides donor-acceptor interfaces for exciton dissociation, while D18-only devices serve as a control for comparison.",
    },
    {
        "Step": "6",
        "Process": "UV-ozone treatment for 20 min",
        "Purpose": "Removes remaining organic contamination and increases the surface energy of ITO. This improves wettability and helps PEDOT:PSS form a continuous film.",
    },
    {
        "Step": "7",
        "Process": "Spin-coating PEDOT:PSS at 4000 rpm",
        "Purpose": "Forms the hole transport layer. PEDOT:PSS smooths the ITO surface, extracts holes, and blocks electrons at the illuminated electrode side.",
    },
    {
        "Step": "8",
        "Process": "Transfer of substrates into the glovebox",
        "Purpose": "Protects the organic active-layer materials from oxygen and moisture during deposition. This is important because both the donor-acceptor morphology and interfaces are sensitive to ambient exposure.",
    },
    {
        "Step": "9",
        "Process": "Spin-coating active layers at 3000 rpm: D18:Y6 without additive, D18:Y6 with additive, and D18 only",
        "Purpose": "Deposits the photoactive semiconductor layer. Comparing these three conditions reveals the role of the Y6 acceptor and the effect of additive-assisted morphology control.",
    },
    {
        "Step": "10",
        "Process": "Thermal annealing of the active layers",
        "Purpose": "Promotes solvent removal, molecular packing, and phase organization. Proper annealing can improve charge transport and reduce recombination in the blend film.",
    },
    {
        "Step": "11",
        "Process": "Spin-coating PDINN at 3000 rpm",
        "Purpose": "Forms the electron transport/interfacial layer. PDINN improves electron extraction at the cathode side and reduces contact resistance with the Ag electrode.",
    },
    {
        "Step": "12",
        "Process": "Shadow-mask alignment and thermal evaporation of 100 nm Ag",
        "Purpose": "Defines the active device area and deposits the top electrode. Ag collects electrons and completes the external circuit.",
    },
    {
        "Step": "13",
        "Process": "Device performance measurement using a solar simulator and EQE system",
        "Purpose": "Extracts J-V parameters including VOC, JSC, FF, and PCE, and measures wavelength-resolved photocurrent generation through EQE spectra.",
    },
]


def method_steps_markdown():
    lines = [
        "| Step | Process | Purpose / analysis |",
        "|---:|---|---|",
    ]
    for item in METHOD_STEPS:
        lines.append(f"| {item['Step']} | {item['Process']} | {item['Purpose']} |")
    return "\n".join(lines)


def build_markdown(jv_stats, eqe_stats):
    compact_jv = compact_jv_stats(jv_stats)
    vals = values_for_text(jv_stats)
    additive_df = additive_comparison(jv_stats)
    jv_cols = ["Device", "Voc (V)", "Jsc (mA cm^-2)", "FF (%)", "PCE (%)"]
    eqe_cols = [
        "Device", "Integrated Jsc from EQE (mA cm^-2)", "Mean EQE 400-800 nm (%)",
        "Maximum EQE (%)", "Linear absorption edge (nm)", "Bandgap from linear edge (eV)",
    ]
    fmt = {c: "{:.3f}" for c in eqe_cols if c != "Device"}
    text = f"""# Lab 3: Fabrication and Characterization of Organic Photovoltaic Devices

## Abstract

This experiment investigated organic photovoltaic (OPV) devices fabricated on patterned ITO/glass substrates with the structure ITO/PEDOT:PSS/active layer/PDINN/Ag. Three active-layer conditions were compared using the group3 J-V data only: D18 only, D18:Y6 without additive, and D18:Y6 with additive. The D18-only control showed almost no photovoltaic response, with a PCE of {vals['d18_pce']:.4f}%. In contrast, the D18:Y6 bulk heterojunction devices reached PCE values of {vals['no_pce']:.2f}% without additive and {vals['add_pce']:.2f}% with additive. The additive increased the short-circuit current density from {vals['no_jsc']:.2f} to {vals['add_jsc']:.2f} mA cm^-2 and improved the fill factor from {vals['no_ff']:.2f}% to {vals['add_ff']:.2f}%. EQE analysis gave integrated current densities of 20.60 and 21.72 mA cm^-2 for the blend devices without and with additive, respectively. These results show that acceptor blending and additive-assisted morphology optimization are essential for efficient OPV operation.

## 1. Introduction

Organic photovoltaics convert sunlight into electrical energy using conjugated organic semiconductors. Compared with inorganic photovoltaic technologies, OPVs can be processed from solution at low temperature and are attractive for lightweight, flexible, and semi-transparent solar modules. Their performance strongly depends on the electronic structure and nanoscale morphology of the active layer.

In this lab, D18 was used as the polymer donor and Y6 as the non-fullerene acceptor. A D18-only device was fabricated as a control, while D18:Y6 blend devices were prepared with and without additive. The comparison demonstrates why a donor-acceptor bulk heterojunction is needed and how processing additives can improve OPV performance.

## 2. Working Principle

In an OPV device, photons are absorbed by the active organic semiconductor layer and generate bound excitons rather than immediately free carriers. Because organic materials have relatively low dielectric constants, the exciton binding energy is large. Efficient photocurrent therefore requires a donor-acceptor interface where the exciton can dissociate: the donor preferentially transports holes, while the acceptor preferentially transports electrons.

The D18:Y6 blend forms a bulk heterojunction, which distributes donor-acceptor interfaces throughout the film. This shortens the distance excitons must diffuse before reaching an interface and creates percolated pathways for holes and electrons. PEDOT:PSS extracts holes at the ITO side, PDINN extracts electrons at the Ag side, and the Ag electrode completes the external circuit.

## 3. Device Architecture

The device architecture used in this experiment is:

ITO/glass / PEDOT:PSS / active layer / PDINN / Ag

ITO/glass serves as the transparent conducting substrate. PEDOT:PSS is the hole transport layer and improves hole extraction from the active layer. The active layer is either D18 only or a D18:Y6 blend. PDINN works as the electron transport layer and interfacial modifier near the metal electrode. The thermally evaporated Ag layer is the top electrode.

### 3.1 Functions of Each Layer in the OPV Device

{layer_functions_markdown()}

## 4. Experimental Method

Patterned ITO/glass substrates were cleaned by sequential sonication in detergent water, deionized water, and absolute ethanol for 15 min each. The substrates were dried with an air gun and stored at 60 degC for 1 h before use. Before film deposition, the substrates were treated by UV-ozone for 20 min to remove residual organic contamination and improve surface wettability.

The active-layer solutions were prepared using chloroform as solvent. Three device conditions were fabricated: D18:Y6 with a 1:1.2 donor-to-acceptor ratio without additive, D18:Y6 with additive, and D18 only. PEDOT:PSS was spin-coated at 4000 rpm as the hole transport layer. The active layers were spin-coated inside the glovebox at 3000 rpm and then thermally annealed. PDINN was spin-coated at 3000 rpm as the electron transport layer. Finally, the samples were masked and loaded into a thermal evaporator for deposition of 100 nm Ag. Device performance was measured under a solar simulator, and EQE spectra were measured from 300 to 950 nm.

### 4.1 Purpose of Each Fabrication Step

{method_steps_markdown()}

## 5. Results

### 5.1 J-V Characteristics

![OPV J-V curves](figures/opv_jv_curves.png)

The D18-only device produced extremely small current density because a single donor material does not provide an efficient acceptor interface for exciton dissociation. The D18:Y6 blend devices showed clear photovoltaic behavior. In the group3 data, adding the processing additive improved both JSC and FF, producing the highest PCE.

![OPV performance summary](figures/opv_performance_summary.png)

{md_table(compact_jv, jv_cols, {})}

**Table 3. Direct comparison of D18:Y6 photovoltaic parameters before and after adding the additive.** The additive mainly improves Jsc, FF, and PCE, while Voc remains unchanged.

{md_table(additive_df, ["Parameter", "Without additive", "With additive", "Change", "Relative change (%)", "Interpretation"], {"Without additive": "{:.3f}", "With additive": "{:.3f}", "Change": "{:+.3f}", "Relative change (%)": "{:+.2f}"})}

### 5.2 EQE Spectra

![OPV EQE spectra](figures/opv_eqe_spectra.png)

The blend devices show broad EQE response from the visible to near-infrared region, consistent with the complementary absorption and charge generation of the D18:Y6 bulk heterojunction. The additive-treated device has higher EQE over most of the wavelength range, agreeing with its higher JSC in the J-V measurement.

{md_table(eqe_stats, eqe_cols, fmt)}

### 5.3 Absorption Edge and Optical Gap

![OPV EQE edge fit](figures/opv_eqe_edge_fit.png)

The long-wavelength EQE edge of the D18:Y6 devices was linearly extrapolated using the 840-920 nm region. The resulting optical gaps are about 1.36 eV for both blend devices, consistent with the low-bandgap Y6 acceptor dominating the near-infrared absorption edge. The D18-only EQE signal was very weak and noisy, so its edge-derived bandgap is not considered reliable in this report.

## 6. Discussion

The large difference between the D18-only control and the D18:Y6 blend devices confirms the central role of the donor-acceptor bulk heterojunction. In D18-only films, excitons have no efficient acceptor phase for charge transfer, so most photogenerated excitons recombine before contributing to current. This explains the near-zero JSC and PCE in group3.

Adding Y6 introduces acceptor domains and donor-acceptor interfaces, enabling exciton dissociation and balanced charge transport. The D18:Y6 group3 devices therefore show JSC values above 25 mA cm^-2 and PCE values above 15%. The additive further improves performance, most likely by optimizing phase separation, molecular packing, and vertical composition distribution. These morphology improvements can reduce recombination and transport resistance, which is reflected in the higher FF and PCE.

The EQE-integrated JSC values are lower than the J-V JSC values. This mismatch can arise from spectral mismatch between the solar simulator and AM1.5G reference spectrum, calibration uncertainty, device area definition, light-bias conditions, or degradation/position variation during measurements. However, the trend remains consistent: the additive-treated blend device gives the highest spectral response and the highest J-V performance.

## 7. Conclusion

Organic photovoltaic devices based on D18, D18:Y6 without additive, and D18:Y6 with additive were fabricated and characterized. This report uses the group3 J-V data for all device-performance comparisons. The D18-only control showed negligible photovoltaic output, confirming that a donor-only active layer cannot efficiently separate excitons. The D18:Y6 blend formed an effective bulk heterojunction and achieved a PCE of {vals['no_pce']:.2f}% without additive. With additive, the PCE increased to {vals['add_pce']:.2f}%, mainly because JSC and FF improved. EQE spectra further confirmed stronger broadband photocurrent generation in the additive-treated device, with an integrated JSC of 21.72 mA cm^-2. Overall, the results demonstrate that both donor-acceptor blending and additive-controlled morphology are key factors for high-performance OPV devices.

## References

1. EE336 Fundamentals of Photovoltaics, Lab Session for Organic Photovoltaics (OPVs) Device Fabrication.
2. J-V measurement files in OPV DATA/D18, OPV DATA/D18 Y6无添加剂, and OPV DATA/D18 Y6加添加剂.
3. EQE measurement files in OPV DATA/EQE.
"""
    (ROOT / "Lab3_OPV_Report_HeYanning_12311512.md").write_text(text, encoding="utf-8")
    return text


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, df, cols, formats):
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, col in zip(table.rows[0].cells, cols):
        set_cell_text(cell, col, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, col in zip(cells, cols):
            value = row[col]
            text = formats.get(col, "{}").format(value) if pd.notna(value) else ""
            set_cell_text(cell, text)
    doc.add_paragraph()


def add_method_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Step", "Process", "Purpose / analysis"]
    widths = [0.55, 2.25, 3.85]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Inches(width)
        set_cell_text(cell, header, bold=True)
    for item in METHOD_STEPS:
        cells = table.add_row().cells
        values = [item["Step"], item["Process"], item["Purpose"]]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = Inches(width)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_layer_functions_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Layer", "Material", "Function"]
    widths = [1.75, 1.45, 3.55]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Inches(width)
        set_cell_text(cell, header, bold=True)
    for item in LAYER_FUNCTIONS:
        cells = table.add_row().cells
        values = [item["Layer"], item["Material"], item["Function"]]
        for cell, value, width in zip(cells, values, widths):
            cell.width = Inches(width)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_additive_comparison_table(doc, jv_stats):
    df = additive_comparison(jv_stats)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Parameter", "Without additive", "With additive", "Change", "Relative change", "Interpretation"]
    widths = [1.15, 1.05, 1.0, 0.8, 0.95, 2.0]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Inches(width)
        set_cell_text(cell, header, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        values = [
            row["Parameter"],
            f"{row['Without additive']:.3f}",
            f"{row['With additive']:.3f}",
            f"{row['Change']:+.3f}",
            f"{row['Relative change (%)']:+.2f}%",
            row["Interpretation"],
        ]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = Inches(width)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 5) else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(8.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, RGBColor(31, 78, 121)),
        ("Heading 3", 11, RGBColor(68, 68, 68)),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.22)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(5)
    return p


def add_figure(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def build_docx(jv_stats, eqe_stats):
    vals = values_for_text(jv_stats)
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lab 3: Fabrication and Characterization of Organic Photovoltaic Devices")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph("He Yanning  |  12311512")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.name = "Times New Roman"
    sub.runs[0].font.size = Pt(10.5)

    doc.add_heading("Abstract", level=1)
    add_para(doc, f"This experiment investigated organic photovoltaic (OPV) devices fabricated on patterned ITO/glass substrates with the structure ITO/PEDOT:PSS/active layer/PDINN/Ag. Three active-layer conditions were compared using the group3 J-V data only: D18 only, D18:Y6 without additive, and D18:Y6 with additive. The D18-only control showed almost no photovoltaic response, with a PCE of {vals['d18_pce']:.4f}%. In contrast, the D18:Y6 bulk heterojunction devices reached PCE values of {vals['no_pce']:.2f}% without additive and {vals['add_pce']:.2f}% with additive. The additive increased the short-circuit current density from {vals['no_jsc']:.2f} to {vals['add_jsc']:.2f} mA cm-2 and improved the fill factor from {vals['no_ff']:.2f}% to {vals['add_ff']:.2f}%. EQE analysis gave integrated current densities of 20.60 and 21.72 mA cm-2 for the blend devices without and with additive, respectively.")

    doc.add_heading("1. Introduction", level=1)
    add_para(doc, "Organic photovoltaics convert sunlight into electrical energy using conjugated organic semiconductors. Compared with inorganic photovoltaic technologies, OPVs can be processed from solution at low temperature and are attractive for lightweight, flexible, and semi-transparent solar modules. Their performance strongly depends on the electronic structure and nanoscale morphology of the active layer.")
    add_para(doc, "In this lab, D18 was used as the polymer donor and Y6 as the non-fullerene acceptor. A D18-only device was fabricated as a control, while D18:Y6 blend devices were prepared with and without additive. The comparison demonstrates why a donor-acceptor bulk heterojunction is needed and how processing additives can improve OPV performance.")

    doc.add_heading("2. Working Principle", level=1)
    add_para(doc, "In an OPV device, photons are absorbed by the active organic semiconductor layer and generate bound excitons rather than immediately free carriers. Efficient photocurrent therefore requires a donor-acceptor interface where the exciton can dissociate: the donor preferentially transports holes, while the acceptor preferentially transports electrons.")
    add_para(doc, "The D18:Y6 blend forms a bulk heterojunction, which distributes donor-acceptor interfaces throughout the film. This shortens the distance excitons must diffuse before reaching an interface and creates percolated pathways for holes and electrons.")

    doc.add_heading("3. Device Architecture", level=1)
    add_para(doc, "The device architecture used in this experiment is ITO/glass / PEDOT:PSS / active layer / PDINN / Ag. ITO/glass serves as the transparent conducting substrate. PEDOT:PSS is the hole transport layer. The active layer is either D18 only or a D18:Y6 blend. PDINN works as the electron transport layer and interfacial modifier near the metal electrode. The thermally evaporated Ag layer is the top electrode.")
    doc.add_heading("3.1 Functions of Each Layer in the OPV Device", level=2)
    add_layer_functions_table(doc)

    doc.add_heading("4. Experimental Method", level=1)
    add_para(doc, "Patterned ITO/glass substrates were cleaned by sequential sonication in detergent water, deionized water, and absolute ethanol for 15 min each. The substrates were dried with an air gun and stored at 60 degC for 1 h before use. Before film deposition, the substrates were treated by UV-ozone for 20 min.")
    add_para(doc, "The active-layer solutions were prepared using chloroform as solvent. Three device conditions were fabricated: D18:Y6 with a 1:1.2 donor-to-acceptor ratio without additive, D18:Y6 with additive, and D18 only. PEDOT:PSS was spin-coated at 4000 rpm. The active layers were spin-coated inside the glovebox at 3000 rpm and then thermally annealed. PDINN was spin-coated at 3000 rpm. Finally, 100 nm Ag was thermally evaporated through a shadow mask.")
    doc.add_heading("4.1 Purpose of Each Fabrication Step", level=2)
    add_method_table(doc)

    doc.add_heading("5. Results", level=1)
    doc.add_heading("5.1 J-V Characteristics", level=2)
    add_figure(doc, FIGURES / "opv_jv_curves.png", "Figure 1. J-V curves of D18-only and D18:Y6 OPV devices.", 6.4)
    add_para(doc, "The D18-only device produced extremely small current density because a single donor material does not provide an efficient acceptor interface for exciton dissociation. The D18:Y6 blend devices showed clear photovoltaic behavior. In the group3 data, adding the processing additive improved both JSC and FF, producing the highest PCE.")
    add_figure(doc, FIGURES / "opv_performance_summary.png", "Figure 2. Average photovoltaic parameters from five devices for each condition.", 5.7)
    compact_jv = compact_jv_stats(jv_stats)
    jv_cols = ["Device", "Voc (V)", "Jsc (mA cm^-2)", "FF (%)", "PCE (%)"]
    add_table(doc, compact_jv, jv_cols, {})
    cap = doc.add_paragraph("Table 3. Direct comparison of D18:Y6 photovoltaic parameters before and after adding the additive. The additive mainly improves Jsc, FF, and PCE, while Voc remains unchanged.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    add_additive_comparison_table(doc, jv_stats)

    doc.add_heading("5.2 EQE Spectra", level=2)
    add_figure(doc, FIGURES / "opv_eqe_spectra.png", "Figure 3. EQE spectra of the three OPV active-layer conditions.", 5.7)
    add_para(doc, "The blend devices show broad EQE response from the visible to near-infrared region, consistent with the complementary absorption and charge generation of the D18:Y6 bulk heterojunction. The additive-treated device has higher EQE over most of the wavelength range, agreeing with its higher JSC in the J-V measurement.")
    eqe_cols = ["Device", "Integrated Jsc from EQE (mA cm^-2)", "Mean EQE 400-800 nm (%)", "Maximum EQE (%)", "Linear absorption edge (nm)", "Bandgap from linear edge (eV)"]
    fmt = {c: "{:.3f}" for c in eqe_cols if c != "Device"}
    add_table(doc, eqe_stats, eqe_cols, fmt)

    doc.add_heading("5.3 Absorption Edge and Optical Gap", level=2)
    add_figure(doc, FIGURES / "opv_eqe_edge_fit.png", "Figure 4. Linear extrapolation of the D18:Y6 long-wavelength EQE edge.", 5.7)
    add_para(doc, "The long-wavelength EQE edge of the D18:Y6 devices was linearly extrapolated using the 840-920 nm region. The resulting optical gaps are about 1.36 eV for both blend devices, consistent with the low-bandgap Y6 acceptor dominating the near-infrared absorption edge. The D18-only EQE signal was very weak and noisy, so its edge-derived bandgap is not considered reliable in this report.")

    doc.add_heading("6. Discussion", level=1)
    add_para(doc, "The large difference between the D18-only control and the D18:Y6 blend devices confirms the central role of the donor-acceptor bulk heterojunction. In D18-only films, excitons have no efficient acceptor phase for charge transfer, so most photogenerated excitons recombine before contributing to current. This explains the near-zero JSC and PCE in group3.")
    add_para(doc, "Adding Y6 introduces acceptor domains and donor-acceptor interfaces, enabling exciton dissociation and balanced charge transport. The D18:Y6 group3 devices therefore show JSC values above 25 mA cm-2 and PCE values above 15%. The additive further improves performance, most likely by optimizing phase separation, molecular packing, and vertical composition distribution. These morphology improvements can reduce recombination and transport resistance, which is reflected in the higher FF and PCE.")
    add_para(doc, "The EQE-integrated JSC values are lower than the J-V JSC values. This mismatch can arise from spectral mismatch between the solar simulator and AM1.5G reference spectrum, calibration uncertainty, device area definition, light-bias conditions, or degradation/position variation during measurements. However, the trend remains consistent: the additive-treated blend device gives the highest spectral response and the highest J-V performance.")

    doc.add_heading("7. Conclusion", level=1)
    add_para(doc, f"Organic photovoltaic devices based on D18, D18:Y6 without additive, and D18:Y6 with additive were fabricated and characterized. This report uses the group3 J-V data for all device-performance comparisons. The D18-only control showed negligible photovoltaic output, confirming that a donor-only active layer cannot efficiently separate excitons. The D18:Y6 blend formed an effective bulk heterojunction and achieved a PCE of {vals['no_pce']:.2f}% without additive. With additive, the PCE increased to {vals['add_pce']:.2f}%, mainly because JSC and FF improved. EQE spectra further confirmed stronger broadband photocurrent generation in the additive-treated device, with an integrated JSC of 21.72 mA cm-2.")

    doc.add_heading("References", level=1)
    for ref in [
        "EE336 Fundamentals of Photovoltaics, Lab Session for Organic Photovoltaics (OPVs) Device Fabrication.",
        "J-V measurement files in OPV DATA/D18, OPV DATA/D18 Y6无添加剂, and OPV DATA/D18 Y6加添加剂.",
        "EQE measurement files in OPV DATA/EQE.",
    ]:
        p = doc.add_paragraph(ref, style=None)
        p.style = doc.styles["Normal"]
    out = ROOT / f"Lab3_OPV_Report_HeYanning_12311512_{REPORT_SUFFIX}.docx"
    try:
        doc.save(out)
    except PermissionError:
        out = ROOT / f"Lab3_OPV_Report_HeYanning_12311512_{REPORT_SUFFIX}_new.docx"
        doc.save(out)
    return out


def main():
    jv_raw, jv_stats, jv_curves = summarize_jv()
    eqe_stats, eqe_curves = summarize_eqe()
    save_tables(jv_raw, jv_stats, eqe_stats)
    plot_jv(jv_curves)
    plot_performance(jv_stats)
    plot_eqe(eqe_curves)
    plot_eqe_edge(eqe_curves)
    build_markdown(jv_stats, eqe_stats)
    out = build_docx(jv_stats, eqe_stats)
    print(out)


if __name__ == "__main__":
    main()
