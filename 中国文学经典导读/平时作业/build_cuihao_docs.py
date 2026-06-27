from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SPEECH_PATH = ROOT / "姓名 学号 平时作业 崔颢与黄鹤楼——一首诗里的武汉 演讲稿.docx"
ARTICLE_PATH = ROOT / "姓名 学号 平时作业 崔颢与黄鹤楼——一首诗里的武汉 公众号推文.docx"


TITLE = "崔颢与黄鹤楼：一首诗里的武汉"
SUBTITLE = "中国文学经典导读平时作业"

POEM = [
    "昔人已乘黄鹤去，此地空余黄鹤楼。",
    "黄鹤一去不复返，白云千载空悠悠。",
    "晴川历历汉阳树，芳草萋萋鹦鹉洲。",
    "日暮乡关何处是？烟波江上使人愁。",
]

SPEECH_PARAGRAPHS = [
    "大家好，今天我想介绍的，是一位和武汉关系非常密切的古代作家：唐代诗人崔颢。严格地说，崔颢并不是武汉本地人，但他写下的《黄鹤楼》，却几乎成了武汉最重要的文学名片之一。很多人来到武汉，未必先背得出武汉的历史沿革，却往往能先说出“昔人已乘黄鹤去，此地空余黄鹤楼”。从这个意义上说，崔颢虽然身在唐代，但他的诗直到今天仍然参与着武汉这座城市的文化表达。",
    "先简单介绍一下崔颢。崔颢是盛唐时期的诗人，擅长写景，也善于把个人情绪融入空间感很强的诗歌中。他最有名的作品就是《黄鹤楼》。这首诗表面上写登楼所见，实际上却把历史传说、眼前景物和人的乡愁放在了一起。前两联从仙人乘鹤的传说写起，写黄鹤已去，只留下空楼与白云；后两联转到现实景色，汉阳树、鹦鹉洲清晰可见，但在这样开阔的景象中，诗人最后感到的却不是豪迈，而是“日暮乡关何处是”的惆怅。",
    "这首诗为什么能成为经典？我觉得至少有三点。第一，它把神话传说和现实景物结合得非常自然，让黄鹤楼不只是一个建筑，而是一个有历史纵深、有想象空间的文化地标。第二，它的景物描写层次极强，从“黄鹤”“白云”写到“汉阳树”“鹦鹉洲”，视野不断展开，读者很容易在脑海里形成画面。第三，它最终落到乡愁，这种情感非常普遍，也因此更容易跨越时代，引起后人的共鸣。",
    "按照课程要求，作业还需要谈到相关纪念活动。崔颢与武汉的关系，最直接的纪念空间当然就是黄鹤楼。今天的黄鹤楼景区，不只是旅游景点，也是诗词文化传播的重要现场。游客进入黄鹤楼，最常接触到的文学内容之一就是《黄鹤楼》；景区内长期以诗词、碑刻、展陈、讲解等方式强化这首诗与城市记忆之间的联系。除了主楼，武汉的晴川阁也和崔颢有关，因为“晴川历历汉阳树”这句诗，已经成为武汉山水意象中极具代表性的表达。",
    "近些年，武汉还不断把黄鹤楼诗词文化转化为公众活动。比如面向社会的黄鹤楼诗词邀请赛、景区中的诗词诵读活动、与夜游演艺结合的文化体验等，都说明《黄鹤楼》并没有停留在课本里，而是进入了今天的城市生活。它既是文学经典，也是市民和游客共同参与的文化资源。这一点很重要，因为它让我们看到，文学经典不是静止的文献，而是可以在当代被重新激活的文化记忆。",
    "最后说一说我自己的理解。我觉得崔颢最值得我们重视的，不只是他写了一首名诗，而是他让一座城市拥有了可以反复讲述、反复感受的精神空间。武汉的黄鹤楼之所以特别，不只是因为楼高景美，更因为它被诗歌赋予了时间感和情感深度。读《黄鹤楼》，我们会意识到，中国文学经典常常不是抽象的，它会落在山川、楼阁、江水、渡口这些具体事物上，然后把一个地方变成文化记忆的容器。",
    "所以，在我看来，崔颢与黄鹤楼的意义，正好体现了文学经典和中国文化之间的关系：经典作品塑造地方意象，地方意象反过来保存经典记忆，二者在历史中不断互相成全。今天我们再读《黄鹤楼》，读到的既是一个唐代诗人的乡愁，也是武汉这座城市延续至今的文化气质。我的分享就到这里，谢谢大家。",
]

ARTICLE_SECTIONS = [
    (
        "一、为什么是崔颢？",
        [
            "如果要在“武汉相关的古代作家”中选一位最有代表性的人物，崔颢几乎是一个很自然的答案。原因不复杂：他虽然不是武汉本地人，却用一首《黄鹤楼》把自己永久写进了武汉的城市记忆。提到黄鹤楼，人们首先想到的往往不是建筑年代，而是“昔人已乘黄鹤去，此地空余黄鹤楼”这两句诗。某种意义上，崔颢用文学完成了一次对城市的命名。",
            "这也非常符合《中国文学经典导读》这门课的思路。文学经典并不是只属于纸上的文本，它还会进入现实生活，进入一座城的风景、气质与共同记忆之中。崔颢与武汉的关系，就是一个很典型的例子。",
        ],
    ),
    (
        "二、作家与作品概述：一首诗如何写出千年名楼",
        [
            "崔颢是唐代诗人，生活在盛唐时期。他最著名的作品就是《黄鹤楼》：",
            "昔人已乘黄鹤去，此地空余黄鹤楼。\n黄鹤一去不复返，白云千载空悠悠。\n晴川历历汉阳树，芳草萋萋鹦鹉洲。\n日暮乡关何处是？烟波江上使人愁。",
            "这首诗最动人的地方，在于它把传说、风景与情感写成了一个完整的精神空间。前两联从黄鹤仙人的传说落笔，让黄鹤楼一开始就带有超越现实的历史感；中间两联又把视线拉回到眼前景物，汉阳树与鹦鹉洲清晰可见，画面开阔而明净；结尾却突然落到“使人愁”，把浩渺江景转化成深沉乡愁。这样一来，黄鹤楼就不再只是楼，而成为寄托时间感、历史感与人生感慨的文化意象。",
            "从文学史角度看，这首诗的经典性至少体现在三点：一是意象完整，黄鹤、白云、汉阳树、鹦鹉洲共同构成极强的空间画面；二是情感自然，乡愁不是硬加上去的，而是在景物展开后顺势生成；三是影响深远，它不仅流传极广，还直接参与了后世对武汉的想象方式。",
        ],
    ),
    (
        "三、武汉的相关纪念活动：诗歌怎样留在城市里",
        [
            "按照课程要求，平时作业不仅要介绍作家和作品，还要谈相关纪念活动。崔颢的纪念方式非常特别，它并不主要表现为单独的“纪念馆”，而是体现在武汉整座黄鹤楼文化空间中。",
            "首先，黄鹤楼景区本身就是崔颢诗意最集中的承载地。游客登楼、观景、听讲解、看碑刻与展陈时，最容易接触到的文学文本之一就是《黄鹤楼》。这首诗已经成为黄鹤楼文化传播的核心内容之一。",
            "其次，崔颢的诗句已经进入武汉的地标体系。比如“晴川历历汉阳树”中的“晴川”，正是后来晴川阁这一城市文化地标的重要命名来源之一。也就是说，崔颢不仅写了黄鹤楼，还参与塑造了武汉城市空间的诗意表达。",
            "再次，武汉近年的文化活动也不断激活黄鹤楼诗词传统。无论是黄鹤楼诗词邀请赛、背诗入园、景区诗词互动，还是把古诗词融入夜游演艺和公众文化体验，都说明经典诗歌并没有停留在教材和注释里，而是在现代城市文化中继续“活着”。",
        ],
    ),
    (
        "四、我的理解：文学经典怎样塑造城市气质",
        [
            "在我看来，崔颢与武汉的关系，最值得注意的地方在于：文学经典能够塑造一座城市的精神气质。黄鹤楼之所以成为武汉的文化象征，不仅因为它历史悠久，更因为历代诗文不断赋予它意义，而崔颢的《黄鹤楼》恰恰是其中最有代表性的一首。",
            "这首诗把“空间”变成了“文化”。当人们看到黄鹤楼，不只是看到一座楼，也会想到黄鹤传说、白云悠悠、江天暮色与游子乡愁。也就是说，诗歌让自然和建筑拥有了可传承的情感结构。一个地方一旦拥有这种结构，就会从地理意义上的“地点”变成文化意义上的“名胜”。",
            "这也让我更理解“文学经典与中国文化的关系”这一课程命题。中国文学经典往往不是抽象地谈人生，而是通过山水、楼阁、边塞、庭院、江河，把情感和价值观寄托在具体世界中。正因为如此，经典才有可能被世代记忆，也才能不断进入现实生活。",
        ],
    ),
    (
        "五、结语",
        [
            "如果说武汉是一座能够被诗歌照亮的城市，那么崔颢就是其中极重要的一位“命名者”。他没有长期生活在武汉，却通过《黄鹤楼》深刻影响了后人理解武汉的方式。今天我们重读这首诗，不只是为了背诵名句，更是为了理解：文学经典为什么能够跨越千年，仍然参与我们对一座城市、一个民族文化传统的认识。",
            "也许这正是经典最有力量的地方。它既属于过去，也始终在场。",
        ],
    ),
]

REFERENCES = [
    "武汉市人民政府门户网站：《官宣！背三首诗就能免费游黄鹤楼公园》",
    "武汉市人民政府门户网站：《黄鹤楼文化展示馆今起开放》",
    "武汉市人民政府门户网站：《晴川阁恢复开放首日，全国游客沉浸体验武汉诗意春天》",
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_normal(doc: Document, font_name: str, size: int, after: int) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(size)
    normal.font.color.rgb = RGBColor(0x18, 0x18, 0x18)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    if name in doc.styles:
        return doc.styles[name]
    return doc.styles.add_style(name, style_type)


def configure_google_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    style_normal(doc, "Arial", 11, 8)

    title_style = ensure_style(doc, "DocTitle")
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_style.font.size = Pt(26)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.font.bold = False
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(3)

    subtitle_style = ensure_style(doc, "DocSubtitle")
    subtitle_style.font.name = "Arial"
    subtitle_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(10)

    for name, size, after, color in [
        ("Heading 1", 20, 6, RGBColor(0, 0, 0)),
        ("Heading 2", 16, 6, RGBColor(0, 0, 0)),
        ("Heading 3", 14, 4, RGBColor(0x43, 0x43, 0x43)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(18 if name == "Heading 2" else 20 if name == "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(after)

    quote_style = ensure_style(doc, "QuoteBlock")
    quote_style.font.name = "Arial"
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    quote_style.font.size = Pt(12)
    quote_style.font.italic = False
    quote_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    quote_style.paragraph_format.left_indent = Inches(0.35)
    quote_style.paragraph_format.right_indent = Inches(0.15)
    quote_style.paragraph_format.space_before = Pt(6)
    quote_style.paragraph_format.space_after = Pt(6)
    quote_style.paragraph_format.line_spacing = 1.2


def configure_business_brief(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style_normal(doc, "Calibri", 11, 6)

    title_style = ensure_style(doc, "SpeechTitle")
    title_style.font.name = "Calibri"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x19, 0x3A, 0x5A)
    title_style.paragraph_format.space_after = Pt(6)

    meta_style = ensure_style(doc, "SpeechMeta")
    meta_style.font.name = "Calibri"
    meta_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    meta_style.font.size = Pt(10.5)
    meta_style.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    meta_style.paragraph_format.space_after = Pt(10)

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, subtitle: str, title_style: str, subtitle_style: str) -> None:
    p = doc.add_paragraph(style=title_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_east_asia_font(run, "微软雅黑")

    sub = doc.add_paragraph(style=subtitle_style)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    set_east_asia_font(run, "微软雅黑")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        set_paragraph_spacing(p, after=4)
        run = p.add_run(f"• {item}")
        set_east_asia_font(run, "微软雅黑")


def build_speech_doc() -> None:
    doc = Document()
    configure_business_brief(doc)
    add_title(doc, TITLE, SUBTITLE, "SpeechTitle", "SpeechMeta")

    info = doc.add_paragraph(style="SpeechMeta")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("时长建议：5-7分钟 | 结构对应PPT 8页")
    set_east_asia_font(run, "微软雅黑")

    doc.add_paragraph("一、开场与选题缘由", style="Heading 1")
    for text in SPEECH_PARAGRAPHS[:2]:
        p = doc.add_paragraph(text)
        set_paragraph_spacing(p, after=6)

    doc.add_paragraph("二、作品分析与经典意义", style="Heading 1")
    poem_p = doc.add_paragraph(style="Normal")
    poem_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(poem_p, before=4, after=8, line=1.2)
    run = poem_p.add_run("\n".join(POEM))
    set_east_asia_font(run, "仿宋")
    run.font.size = Pt(12)

    for text in SPEECH_PARAGRAPHS[2:4]:
        p = doc.add_paragraph(text)
        set_paragraph_spacing(p, after=6)

    doc.add_paragraph("三、武汉相关纪念活动", style="Heading 1")
    for text in SPEECH_PARAGRAPHS[4:5]:
        p = doc.add_paragraph(text)
        set_paragraph_spacing(p, after=6)
    add_bullets(
        doc,
        [
            "黄鹤楼景区内长期以诗词、讲解、碑刻与展陈强化《黄鹤楼》的文化记忆。",
            "晴川阁因“晴川历历汉阳树”而与崔颢形成稳定的城市诗意关联。",
            "诗词邀请赛、背诗入园、夜游诗词互动等活动，让经典进入今天的公共文化生活。",
        ],
    )

    doc.add_paragraph("四、个人理解与结语", style="Heading 1")
    for text in SPEECH_PARAGRAPHS[5:]:
        p = doc.add_paragraph(text)
        set_paragraph_spacing(p, after=6)

    doc.save(SPEECH_PATH)


def build_article_doc() -> None:
    doc = Document()
    configure_google_doc_style(doc)
    add_title(doc, TITLE, "从一首《黄鹤楼》读懂文学经典如何塑造城市记忆", "DocTitle", "DocSubtitle")

    lead = doc.add_paragraph(style="QuoteBlock")
    lead.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = lead.add_run("黄鹤楼之于武汉，不只是建筑地标，更是被诗歌反复点亮的文化地标。")
    set_east_asia_font(run, "仿宋")

    meta = doc.add_paragraph(style="DocSubtitle")
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta.add_run("中国文学经典导读平时作业 | 主题：武汉相关古代作家")
    set_east_asia_font(run, "微软雅黑")

    for heading, paragraphs in ARTICLE_SECTIONS:
        doc.add_paragraph(heading, style="Heading 1")
        for idx, text in enumerate(paragraphs):
            if heading == "二、作家与作品概述：一首诗如何写出千年名楼" and idx == 1:
                qp = doc.add_paragraph(style="QuoteBlock")
                run = qp.add_run(text)
                set_east_asia_font(run, "仿宋")
                continue
            p = doc.add_paragraph(text)
            set_paragraph_spacing(p, after=8)

    doc.add_paragraph("参考资料", style="Heading 1")
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        set_paragraph_spacing(p, after=4)

    doc.save(ARTICLE_PATH)


if __name__ == "__main__":
    build_speech_doc()
    build_article_doc()
